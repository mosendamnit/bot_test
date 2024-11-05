import os
from docx import Document as DocxDocument
from langchain.schema.document import Document
from document_utils import split_documents

DATA_PATH = "data"

# function to load data file in docx format
def load_word_documents():
    documents = []
    section_counter = 1

    # Load Word documents
    for filename in os.listdir(DATA_PATH):
        if filename.endswith(".docx"):
            file_path = os.path.join(DATA_PATH, filename)
            doc = DocxDocument(file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Add section number to metadata to simulate pages
            metadata = {"source": filename, "page": f"Section {section_counter}"}
            section_counter += 1

            documents.append(Document(page_content=content, metadata =metadata))
            print(f"Loaded Word Document: {filename}")

    return documents

if __name__ == "__main__":
    
    # Load Word documents
    documents = load_word_documents()
    print(f"Total Word Documents Loaded: {len(documents)}")

    # Use imported function to split documents
    chunks = split_documents(documents)
    print(f"Total Chunks Created: {len(chunks)}")
